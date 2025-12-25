import streamlit as st
import docx
import pdfplumber
import requests
import os

st.title("📄 Insurance GLR Automation – AI Powered")

OPENROUTER_API_KEY = st.secrets["OPENROUTER_KEY"]

def call_llm(prompt):
    url = "https://openrouter.ai/api/v1/chat/completions"
    headers = {"Authorization": f"Bearer {OPENROUTER_API_KEY}"}

    payload = {
        "model": "deepseek/deepseek-chat",
        "messages": [{"role": "user", "content": prompt}]
    }

    response = requests.post(url, json=payload, headers=headers).json()
    return response["choices"][0]["message"]["content"]

def fill_template(template_docx, extracted_text):
    doc = docx.Document(template_docx)

    prompt = f"""
Extract key-value pairs needed to fill this insurance template.
Template content:
{[p.text for p in doc.paragraphs]}

Report extracted text:
{extracted_text}

Return ONLY in 'Field: Value' format.
"""

    result = call_llm(prompt)
    kv_pairs = {}

    for line in result.split("\n"):
        if ":" in line:
            key, value = line.split(":", 1)
            kv_pairs[key.strip()] = value.strip()

    # Replace template fields
    for para in doc.paragraphs:
        for key, val in kv_pairs.items():
            if key in para.text:
                para.text = para.text.replace(key, val)

    output_path = "task_3_output/output.docx"
    doc.save(output_path)
    return output_path

# Streamlit UI -----------------------
template = st.file_uploader("Upload Insurance Template (.docx)", type=["docx"])
reports = st.file_uploader("Upload Photo Reports (.pdf)", type=["pdf"], accept_multiple_files=True)

if st.button("Generate Filled Template"):
    if template and reports:
        extracted = ""

        for pdf_file in reports:
            with pdfplumber.open(pdf_file) as pdf:
                for page in pdf.pages:
                    extracted += page.extract_text() + "\n"

        output = fill_template(template, extracted)
        st.success("Template Generated Successfully!")

        with open(output, "rb") as f:
            st.download_button("Download Filled Template", f, file_name="filled_template.docx")
